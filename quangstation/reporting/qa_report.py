#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Module tạo báo cáo kiểm tra chất lượng kế hoạch (QA) cho QuangStation V2.
Hỗ trợ các định dạng báo cáo khác nhau: PDF, HTML, DOCX.
"""

import os
import json
import tempfile
import shutil
from datetime import datetime
from typing import Dict, List, Any, Optional

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

from quangstation.utils.logging import get_logger

logger = get_logger(__name__)

def create_qa_pdf_report(output_path: str, report_data: Dict[str, Any]) -> str:
    """
    Tạo báo cáo QA định dạng PDF
    
    Args:
        output_path: Đường dẫn file báo cáo đầu ra
        report_data: Dữ liệu báo cáo
        
    Returns:
        Đường dẫn file báo cáo đã tạo
    """
    try:
        # Tạo báo cáo HTML trước
        temp_dir = tempfile.mkdtemp()
        html_path = os.path.join(temp_dir, "temp_qa_report.html")
        
        create_qa_html_report(html_path, report_data)
        
        # Chuyển đổi HTML sang PDF
        try:
            import weasyprint
            weasyprint.HTML(html_path).write_pdf(output_path)
            
            # Xóa file tạm
            shutil.rmtree(temp_dir, ignore_errors=True)
            
            return output_path
        except ImportError:
            logger.warning("Thư viện weasyprint không được cài đặt. Sử dụng báo cáo HTML thay thế.")
            # Sao chép file HTML sang vị trí đầu ra
            output_html_path = output_path.replace('.pdf', '.html')
            shutil.copy2(html_path, output_html_path)
            
            # Xóa file tạm
            shutil.rmtree(temp_dir, ignore_errors=True)
            
            return output_html_path
    
    except Exception as e:
        logger.error(f"Lỗi khi tạo báo cáo PDF QA: {str(e)}")
        raise


def create_qa_html_report(output_path: str, report_data: Dict[str, Any]) -> str:
    """
    Tạo báo cáo QA định dạng HTML
    
    Args:
        output_path: Đường dẫn file báo cáo đầu ra
        report_data: Dữ liệu báo cáo
        
    Returns:
        Đường dẫn file báo cáo đã tạo
    """
    try:
        # Thu thập dữ liệu
        patient_info = report_data.get("patient_info", {})
        plan_info = report_data.get("plan_info", {})
        qa_results = report_data.get("qa_results", {})
        title = report_data.get("title", "Báo cáo kiểm tra chất lượng kế hoạch xạ trị")
        included_sections = report_data.get("included_sections", [])
        
        # Tạo template HTML
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{title}</title>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1, h2, h3 {{ color: #2c3e50; }}
                .section {{ margin-bottom: 30px; border: 1px solid #ddd; padding: 15px; border-radius: 5px; }}
                table {{ border-collapse: collapse; width: 100%; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; }}
                th {{ background-color: #f2f2f2; }}
                .pass {{ color: green; font-weight: bold; }}
                .fail {{ color: red; font-weight: bold; }}
                .center {{ text-align: center; }}
                .img-container {{ text-align: center; margin: 20px 0; }}
                img {{ max-width: 90%; }}
                .footer {{ margin-top: 30px; font-size: 0.8em; color: #7f8c8d; text-align: center; }}
            </style>
        </head>
        <body>
            <h1 class="center">{title}</h1>
            <p class="center">Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
        """
        
        # Thêm thông tin bệnh nhân và kế hoạch
        if "patient_info" in included_sections:
            html_content += """
            <div class="section">
                <h2>Thông tin bệnh nhân</h2>
                <table>
                    <tr>
                        <th>Thông tin</th>
                        <th>Giá trị</th>
                    </tr>
            """
            
            for key, value in patient_info.items():
                html_content += f"""
                    <tr>
                        <td>{key}</td>
                        <td>{value}</td>
                    </tr>
                """
            
            html_content += """
                </table>
            </div>
            """
        
        if "plan_info" in included_sections:
            html_content += """
            <div class="section">
                <h2>Thông tin kế hoạch</h2>
                <table>
                    <tr>
                        <th>Thông tin</th>
                        <th>Giá trị</th>
                    </tr>
            """
            
            for key, value in plan_info.items():
                html_content += f"""
                    <tr>
                        <td>{key}</td>
                        <td>{value}</td>
                    </tr>
                """
            
            html_content += """
                </table>
            </div>
            """
        
        # Thêm kết quả QA
        if "qa_results" in included_sections and qa_results:
            html_content += """
            <div class="section">
                <h2>Kết quả kiểm tra QA</h2>
            """
            
            all_passed = qa_results.get("all_passed", False)
            status = "Đạt" if all_passed else "Không đạt"
            status_class = "pass" if all_passed else "fail"
            
            html_content += f"""
                <p>Kết quả tổng thể: <span class="{status_class}">{status}</span></p>
                
                <h3>Chi tiết kiểm tra</h3>
            """
            
            if "results" in qa_results:
                results = qa_results["results"]
                
                # 1. Độ phủ liều
                if "coverage" in results:
                    coverage = results["coverage"]
                    passed = coverage.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    status_class = "pass" if passed else "fail"
                    
                    html_content += f"""
                    <div class="subsection">
                        <h3>1. Độ phủ liều</h3>
                        <p>Kết quả: <span class="{status_class}">{status}</span></p>
                        <p>{coverage.get("comment", "")}</p>
                        
                        <table>
                            <tr>
                                <th>Chỉ số</th>
                                <th>Giá trị</th>
                            </tr>
                    """
                    
                    if "details" in coverage:
                        details = coverage["details"]
                        for key, value in details.items():
                            html_content += f"""
                                <tr>
                                    <td>{key}</td>
                                    <td>{value:.1f}%</td>
                                </tr>
                            """
                    
                    html_content += """
                        </table>
                    </div>
                    """
                
                # 2. Ràng buộc OAR
                if "oar_constraints" in results:
                    oar = results["oar_constraints"]
                    passed = oar.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    status_class = "pass" if passed else "fail"
                    
                    html_content += f"""
                    <div class="subsection">
                        <h3>2. Ràng buộc cơ quan nguy cấp</h3>
                        <p>Kết quả: <span class="{status_class}">{status}</span></p>
                        <p>{oar.get("comment", "")}</p>
                    """
                    
                    if "results_by_organ" in oar:
                        for organ_name, organ_result in oar["results_by_organ"].items():
                            organ_passed = organ_result.get("passed", False)
                            organ_status = "Đạt" if organ_passed else "Không đạt"
                            organ_status_class = "pass" if organ_passed else "fail"
                            
                            html_content += f"""
                            <h4>{organ_name}: <span class="{organ_status_class}">{organ_status}</span></h4>
                            <table>
                                <tr>
                                    <th>Ràng buộc</th>
                                    <th>Giá trị thực tế</th>
                                    <th>Giới hạn</th>
                                    <th>Kết quả</th>
                                </tr>
                            """
                            
                            if "constraint_results" in organ_result:
                                for constraint_name, constraint_result in organ_result["constraint_results"].items():
                                    constraint_passed = constraint_result.get("passed", False)
                                    constraint_status = "Đạt" if constraint_passed else "Không đạt"
                                    constraint_status_class = "pass" if constraint_passed else "fail"
                                    
                                    html_content += f"""
                                        <tr>
                                            <td>{constraint_name}</td>
                                            <td>{constraint_result.get("actual", 0):.1f}</td>
                                            <td>{constraint_result.get("limit", 0):.1f}</td>
                                            <td class="{constraint_status_class}">{constraint_status}</td>
                                        </tr>
                                    """
                            
                            html_content += """
                            </table>
                            """
                    
                    html_content += """
                    </div>
                    """
                
                # 3. Tính phù hợp
                if "conformity" in results:
                    conformity = results["conformity"]
                    passed = conformity.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    status_class = "pass" if passed else "fail"
                    
                    html_content += f"""
                    <div class="subsection">
                        <h3>3. Tính phù hợp (Conformity)</h3>
                        <p>Kết quả: <span class="{status_class}">{status}</span></p>
                        <p>{conformity.get("comment", "")}</p>
                        
                        <table>
                            <tr>
                                <th>Chỉ số</th>
                                <th>Giá trị</th>
                            </tr>
                    """
                    
                    if "details" in conformity:
                        details = conformity["details"]
                        for key, value in details.items():
                            html_content += f"""
                                <tr>
                                    <td>{key}</td>
                                    <td>{value:.3f}</td>
                                </tr>
                            """
                    
                    html_content += """
                        </table>
                    </div>
                    """
                
                # 4. Tính đồng nhất
                if "homogeneity" in results:
                    homogeneity = results["homogeneity"]
                    passed = homogeneity.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    status_class = "pass" if passed else "fail"
                    
                    html_content += f"""
                    <div class="subsection">
                        <h3>4. Tính đồng nhất (Homogeneity)</h3>
                        <p>Kết quả: <span class="{status_class}">{status}</span></p>
                        <p>{homogeneity.get("comment", "")}</p>
                        
                        <table>
                            <tr>
                                <th>Chỉ số</th>
                                <th>Giá trị</th>
                            </tr>
                    """
                    
                    if "details" in homogeneity:
                        details = homogeneity["details"]
                        for key, value in details.items():
                            html_content += f"""
                                <tr>
                                    <td>{key}</td>
                                    <td>{value:.3f}</td>
                                </tr>
                            """
                    
                    html_content += """
                        </table>
                    </div>
                    """
            
            html_content += """
            </div>
            """
        
        # Thêm phần chân trang
        html_content += """
            <div class="footer">
                <p>Báo cáo được tạo bởi QuangStation V2 - Hệ thống lập kế hoạch xạ trị</p>
            </div>
        </body>
        </html>
        """
        
        # Ghi vào file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path
    
    except Exception as e:
        logger.error(f"Lỗi khi tạo báo cáo HTML QA: {str(e)}")
        raise


def create_qa_docx_report(output_path: str, report_data: Dict[str, Any]) -> str:
    """
    Tạo báo cáo QA định dạng DOCX
    
    Args:
        output_path: Đường dẫn file báo cáo đầu ra
        report_data: Dữ liệu báo cáo
        
    Returns:
        Đường dẫn file báo cáo đã tạo
    """
    try:
        # Kiểm tra thư viện python-docx
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("Thư viện python-docx không được cài đặt. Sử dụng báo cáo HTML thay thế.")
            # Tạo báo cáo HTML thay thế
            output_html_path = output_path.replace('.docx', '.html')
            return create_qa_html_report(output_html_path, report_data)
        
        # Thu thập dữ liệu
        patient_info = report_data.get("patient_info", {})
        plan_info = report_data.get("plan_info", {})
        qa_results = report_data.get("qa_results", {})
        title = report_data.get("title", "Báo cáo kiểm tra chất lượng kế hoạch xạ trị")
        included_sections = report_data.get("included_sections", [])
        
        # Tạo document mới
        document = Document()
        
        # Tiêu đề báo cáo
        title_heading = document.add_heading(title, level=0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ngày tạo
        date_paragraph = document.add_paragraph(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm thông tin bệnh nhân
        if "patient_info" in included_sections:
            document.add_heading("Thông tin bệnh nhân", level=1)
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Tiêu đề bảng
            header_cells = table.rows[0].cells
            header_cells[0].text = "Thông tin"
            header_cells[1].text = "Giá trị"
            
            # Dữ liệu bảng
            for key, value in patient_info.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
            
            document.add_paragraph()
        
        # Thêm thông tin kế hoạch
        if "plan_info" in included_sections:
            document.add_heading("Thông tin kế hoạch", level=1)
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Tiêu đề bảng
            header_cells = table.rows[0].cells
            header_cells[0].text = "Thông tin"
            header_cells[1].text = "Giá trị"
            
            # Dữ liệu bảng
            for key, value in plan_info.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
            
            document.add_paragraph()
        
        # Thêm kết quả QA
        if "qa_results" in included_sections and qa_results:
            document.add_heading("Kết quả kiểm tra QA", level=1)
            
            all_passed = qa_results.get("all_passed", False)
            status = "Đạt" if all_passed else "Không đạt"
            
            result_paragraph = document.add_paragraph(f"Kết quả tổng thể: {status}")
            # Đặt màu cho kết quả
            for run in result_paragraph.runs:
                if status == "Đạt":
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Màu xanh lá
                else:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Màu đỏ
            
            if "results" in qa_results:
                results = qa_results["results"]
                
                # 1. Độ phủ liều
                if "coverage" in results:
                    coverage = results["coverage"]
                    passed = coverage.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    
                    document.add_heading("1. Độ phủ liều", level=2)
                    
                    result_paragraph = document.add_paragraph(f"Kết quả: {status}")
                    # Đặt màu cho kết quả
                    for run in result_paragraph.runs:
                        if status == "Đạt":
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Màu xanh lá
                        else:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Màu đỏ
                    
                    document.add_paragraph(coverage.get("comment", ""))
                    
                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    
                    # Tiêu đề bảng
                    header_cells = table.rows[0].cells
                    header_cells[0].text = "Chỉ số"
                    header_cells[1].text = "Giá trị"
                    
                    # Dữ liệu bảng
                    if "details" in coverage:
                        details = coverage["details"]
                        for key, value in details.items():
                            row_cells = table.add_row().cells
                            row_cells[0].text = key
                            row_cells[1].text = f"{value:.1f}%"
                    
                    document.add_paragraph()
                
                # 2. Ràng buộc OAR
                if "oar_constraints" in results:
                    oar = results["oar_constraints"]
                    passed = oar.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    
                    document.add_heading("2. Ràng buộc cơ quan nguy cấp", level=2)
                    
                    result_paragraph = document.add_paragraph(f"Kết quả: {status}")
                    # Đặt màu cho kết quả
                    for run in result_paragraph.runs:
                        if status == "Đạt":
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Màu xanh lá
                        else:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Màu đỏ
                    
                    document.add_paragraph(oar.get("comment", ""))
                    
                    if "results_by_organ" in oar:
                        for organ_name, organ_result in oar["results_by_organ"].items():
                            organ_passed = organ_result.get("passed", False)
                            organ_status = "Đạt" if organ_passed else "Không đạt"
                            
                            document.add_heading(f"{organ_name}: {organ_status}", level=3)
                            
                            table = document.add_table(rows=1, cols=4)
                            table.style = 'Table Grid'
                            
                            # Tiêu đề bảng
                            header_cells = table.rows[0].cells
                            header_cells[0].text = "Ràng buộc"
                            header_cells[1].text = "Giá trị thực tế"
                            header_cells[2].text = "Giới hạn"
                            header_cells[3].text = "Kết quả"
                            
                            # Dữ liệu bảng
                            if "constraint_results" in organ_result:
                                for constraint_name, constraint_result in organ_result["constraint_results"].items():
                                    constraint_passed = constraint_result.get("passed", False)
                                    constraint_status = "Đạt" if constraint_passed else "Không đạt"
                                    
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = constraint_name
                                    row_cells[1].text = f"{constraint_result.get('actual', 0):.1f}"
                                    row_cells[2].text = f"{constraint_result.get('limit', 0):.1f}"
                                    row_cells[3].text = constraint_status
                                    
                                    # Đặt màu cho kết quả
                                    if constraint_status == "Đạt":
                                        row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                                    else:
                                        row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                            
                            document.add_paragraph()
                
                # 3. Tính phù hợp
                if "conformity" in results:
                    conformity = results["conformity"]
                    passed = conformity.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    
                    document.add_heading("3. Tính phù hợp (Conformity)", level=2)
                    
                    result_paragraph = document.add_paragraph(f"Kết quả: {status}")
                    # Đặt màu cho kết quả
                    for run in result_paragraph.runs:
                        if status == "Đạt":
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Màu xanh lá
                        else:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Màu đỏ
                    
                    document.add_paragraph(conformity.get("comment", ""))
                    
                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    
                    # Tiêu đề bảng
                    header_cells = table.rows[0].cells
                    header_cells[0].text = "Chỉ số"
                    header_cells[1].text = "Giá trị"
                    
                    # Dữ liệu bảng
                    if "details" in conformity:
                        details = conformity["details"]
                        for key, value in details.items():
                            row_cells = table.add_row().cells
                            row_cells[0].text = key
                            row_cells[1].text = f"{value:.3f}"
                    
                    document.add_paragraph()
                
                # 4. Tính đồng nhất
                if "homogeneity" in results:
                    homogeneity = results["homogeneity"]
                    passed = homogeneity.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    
                    document.add_heading("4. Tính đồng nhất (Homogeneity)", level=2)
                    
                    result_paragraph = document.add_paragraph(f"Kết quả: {status}")
                    # Đặt màu cho kết quả
                    for run in result_paragraph.runs:
                        if status == "Đạt":
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Màu xanh lá
                        else:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Màu đỏ
                    
                    document.add_paragraph(homogeneity.get("comment", ""))
                    
                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    
                    # Tiêu đề bảng
                    header_cells = table.rows[0].cells
                    header_cells[0].text = "Chỉ số"
                    header_cells[1].text = "Giá trị"
                    
                    # Dữ liệu bảng
                    if "details" in homogeneity:
                        details = homogeneity["details"]
                        for key, value in details.items():
                            row_cells = table.add_row().cells
                            row_cells[0].text = key
                            row_cells[1].text = f"{value:.3f}"
                    
                    document.add_paragraph()
        
        # Thêm chân trang
        footer_paragraph = document.add_paragraph("Báo cáo được tạo bởi QuangStation V2 - Hệ thống lập kế hoạch xạ trị")
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Lưu document
        document.save(output_path)
        
        return output_path
    
    except Exception as e:
        logger.error(f"Lỗi khi tạo báo cáo DOCX QA: {str(e)}")
        raise 