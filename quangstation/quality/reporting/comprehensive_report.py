#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Module tạo báo cáo tổng hợp cho QuangStation V2.
Kết hợp các loại báo cáo khác nhau vào một báo cáo thống nhất.
"""

import os
import json
import tempfile
import shutil
from datetime import datetime
from typing import Dict, List, Any, Optional
import matplotlib.pyplot as plt

from quangstation.core.utils.logging import get_logger

logger = get_logger(__name__)

class ComprehensiveReport:
    """
    Lớp tạo báo cáo tổng hợp cho kế hoạch xạ trị.
    """
    
    def __init__(self, patient_info: Dict[str, Any], plan_info: Dict[str, Any], 
                 structures: Dict[str, Any] = None, dose_data: Dict[str, Any] = None):
        """
        Khởi tạo báo cáo tổng hợp
        
        Args:
            patient_info: Thông tin bệnh nhân
            plan_info: Thông tin kế hoạch
            structures: Dữ liệu cấu trúc
            dose_data: Dữ liệu liều
        """
        self.patient_info = patient_info or {}
        self.plan_info = plan_info or {}
        self.structures = structures or {}
        self.dose_data = dose_data or {}
        self.included_sections = [
            "patient_info", "plan_info", "dvh", "dose_metrics", 
            "qa_results", "kbp_results", "conformity_indices", "homogeneity_indices"
        ]
        self.images = {}  # Danh sách hình ảnh để thêm vào báo cáo
        self.logger = get_logger("ComprehensiveReport")
    
    def generate(self, output_path: str, format: str = "pdf") -> str:
        """
        Tạo báo cáo tổng hợp
        
        Args:
            output_path: Đường dẫn file báo cáo đầu ra
            format: Định dạng file ('pdf', 'html', 'docx', 'json')
            
        Returns:
            Đường dẫn file báo cáo đã tạo
        """
        # Chuẩn bị dữ liệu
        data = {
            "patient_info": self.patient_info,
            "plan_info": self.plan_info,
            "structures": self.structures,
            "dose_data": self.dose_data,
            "included_sections": self.included_sections,
            "generated_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "images": self.images
        }
        
        # Gọi hàm tạo báo cáo
        return create_comprehensive_report(
            output_path=output_path,
            format=format,
            patient_data=self.patient_info,
            plan_data=self.plan_info,
            included_sections=self.included_sections
        )
    
    def add_image(self, image_path: str, title: str = None) -> None:
        """
        Thêm hình ảnh vào báo cáo
        
        Args:
            image_path: Đường dẫn file hình ảnh
            title: Tiêu đề hình ảnh
        """
        if not os.path.exists(image_path):
            self.logger.warning(f"Không tìm thấy file hình ảnh: {image_path}")
            return
            
        title = title or os.path.basename(image_path)
        self.images[title] = image_path
        self.logger.info(f"Đã thêm hình ảnh: {title}")
    
    def generate_pdf(self, output_path: str) -> str:
        """
        Tạo báo cáo dạng PDF
        
        Args:
            output_path: Đường dẫn file báo cáo đầu ra
            
        Returns:
            Đường dẫn file báo cáo đã tạo
        """
        return self.generate(output_path, format="pdf")
    
    def generate_docx(self, output_path: str) -> str:
        """
        Tạo báo cáo dạng DOCX
        
        Args:
            output_path: Đường dẫn file báo cáo đầu ra
            
        Returns:
            Đường dẫn file báo cáo đã tạo
        """
        return self.generate(output_path, format="docx")
    
    def generate_html(self, output_path: str) -> str:
        """
        Tạo báo cáo dạng HTML
        
        Args:
            output_path: Đường dẫn file báo cáo đầu ra
            
        Returns:
            Đường dẫn file báo cáo đã tạo
        """
        return self.generate(output_path, format="html")

def create_comprehensive_report(
    output_path: str,
    format: str,
    standard_report_path: Optional[str] = None,
    qa_report_path: Optional[str] = None,
    kbp_qa_report_path: Optional[str] = None,
    patient_data: Dict[str, Any] = None,
    plan_data: Dict[str, Any] = None,
    included_sections: List[str] = None
) -> str:
    """
    Tạo báo cáo tổng hợp từ các báo cáo khác nhau

    Args:
        output_path: Đường dẫn file báo cáo đầu ra
        format: Định dạng file ('pdf', 'html', 'docx', 'json')
        standard_report_path: Đường dẫn đến báo cáo tiêu chuẩn
        qa_report_path: Đường dẫn đến báo cáo QA
        kbp_qa_report_path: Đường dẫn đến báo cáo KBP QA
        patient_data: Dữ liệu bệnh nhân
        plan_data: Dữ liệu kế hoạch
        included_sections: Các phần cần đưa vào báo cáo
    
    Returns:
        Đường dẫn file báo cáo đã tạo
    """
    try:
        if included_sections is None:
            included_sections = [
                "patient_info", "plan_info", "dvh", "dose_metrics", 
                "qa_results", "kbp_results", "conformity_indices", "homogeneity_indices"
            ]
        
        # Tạo thư mục tạm để xử lý
        temp_dir = tempfile.mkdtemp()
        
        # Dữ liệu tổng hợp cho báo cáo
        comprehensive_data = {
            "patient_info": patient_data or {},
            "plan_info": plan_data or {},
            "title": f"Báo cáo tổng hợp kế hoạch xạ trị - {datetime.now().strftime('%d/%m/%Y')}",
            "timestamp": datetime.now().isoformat(),
            "included_sections": included_sections
        }
        
        # Đọc và kết hợp dữ liệu từ các báo cáo riêng biệt
        # Báo cáo tiêu chuẩn
        if standard_report_path and os.path.exists(standard_report_path):
            if standard_report_path.endswith('.json'):
                with open(standard_report_path, 'r') as f:
                    standard_data = json.load(f)
                    if "dvh_data" in standard_data:
                        comprehensive_data["dvh_data"] = standard_data["dvh_data"]
                    if "dose_metrics" in standard_data:
                        comprehensive_data["dose_metrics"] = standard_data["dose_metrics"]
            else:
                # Nếu không phải JSON, lưu đường dẫn để tham chiếu sau này
                comprehensive_data["standard_report_path"] = standard_report_path
        
        # Báo cáo QA
        if qa_report_path and os.path.exists(qa_report_path):
            if qa_report_path.endswith('.json'):
                with open(qa_report_path, 'r') as f:
                    qa_data = json.load(f)
                    comprehensive_data["qa_results"] = qa_data
            else:
                comprehensive_data["qa_report_path"] = qa_report_path
        
        # Báo cáo KBP QA
        if kbp_qa_report_path and os.path.exists(kbp_qa_report_path):
            if kbp_qa_report_path.endswith('.json'):
                with open(kbp_qa_report_path, 'r') as f:
                    kbp_qa_data = json.load(f)
                    comprehensive_data["kbp_results"] = kbp_qa_data
            else:
                comprehensive_data["kbp_qa_report_path"] = kbp_qa_report_path
        
        # Tạo báo cáo tổng hợp dựa trên định dạng
        result_path = None
        
        if format == 'json':
            # Xuất trực tiếp dưới dạng JSON
            with open(output_path, 'w') as f:
                json.dump(comprehensive_data, f, indent=2)
            result_path = output_path
        
        elif format == 'html':
            result_path = _create_html_comprehensive_report(output_path, comprehensive_data, temp_dir)
        
        elif format == 'pdf':
            result_path = _create_pdf_comprehensive_report(output_path, comprehensive_data, temp_dir)
        
        elif format == 'docx':
            result_path = _create_docx_comprehensive_report(output_path, comprehensive_data, temp_dir)
        
        else:
            raise ValueError(f"Định dạng không được hỗ trợ: {format}")
        
        # Xóa thư mục tạm
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        logger.info(f"Đã tạo báo cáo tổng hợp tại: {result_path}")
        return result_path
    
    except Exception as error:
        logger.error(f"Lỗi khi tạo báo cáo tổng hợp: {str(error)}")
        raise


def _create_html_comprehensive_report(output_path: str, data: Dict[str, Any], temp_dir: str) -> str:
    """Tạo báo cáo tổng hợp HTML"""
    try:
        # Thu thập dữ liệu
        patient_info = data.get("patient_info", {})
        plan_info = data.get("plan_info", {})
        dvh_data = data.get("dvh_data", {})
        dose_metrics = data.get("dose_metrics", {})
        qa_results = data.get("qa_results", {})
        kbp_results = data.get("kbp_results", [])
        included_sections = data.get("included_sections", [])
        
        # Tạo template HTML
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{data.get("title", "Báo cáo tổng hợp kế hoạch xạ trị")}</title>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1, h2, h3 {{ color: #2c3e50; }}
                .section {{ margin-bottom: 30px; border: 1px solid #ddd; padding: 15px; border-radius: 5px; }}
                table {{ border-collapse: collapse; width: 100%; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; }}
                th {{ background-color: #f2f2f2; }}
                .pass {{ color: green; }}
                .fail {{ color: red; }}
                .center {{ text-align: center; }}
                .img-container {{ text-align: center; margin: 20px 0; }}
                img {{ max-width: 90%; }}
                .footer {{ margin-top: 30px; font-size: 0.8em; color: #7f8c8d; text-align: center; }}
            </style>
        </head>
        <body>
            <h1 class="center">{data.get("title", "Báo cáo tổng hợp kế hoạch xạ trị")}</h1>
            <p class="center">Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
        """
        
        # Thêm thông tin bệnh nhân và kế hoạch
        if "patient_info" in included_sections:
            html_content += """
            <div class="section">
                <h2>Thông tin bệnh nhân</h2>
                <table>
                    <tr>
                        <th>Thông tin</th>
                        <th>Giá trị</th>
                    </tr>
            """
            
            for key, value in patient_info.items():
                html_content += f"""
                    <tr>
                        <td>{key}</td>
                        <td>{value}</td>
                    </tr>
                """
            
            html_content += """
                </table>
            </div>
            """
        
        if "plan_info" in included_sections:
            html_content += """
            <div class="section">
                <h2>Thông tin kế hoạch</h2>
                <table>
                    <tr>
                        <th>Thông tin</th>
                        <th>Giá trị</th>
                    </tr>
            """
            
            for key, value in plan_info.items():
                html_content += f"""
                    <tr>
                        <td>{key}</td>
                        <td>{value}</td>
                    </tr>
                """
            
            html_content += """
                </table>
            </div>
            """
        
        # Thêm biểu đồ DVH nếu có
        if "dvh" in included_sections and "standard_report_path" in data:
            standard_report_path = data["standard_report_path"]
            if standard_report_path.endswith('.html'):
                # Trích xuất hình ảnh DVH từ báo cáo HTML tiêu chuẩn
                import re
                with open(standard_report_path, 'r') as f:
                    html_content_standard = f.read()
                    dvh_image_match = re.search(r'<img src="data:image/png;base64,[^"]+"', html_content_standard)
                    if dvh_image_match:
                        dvh_image_tag = dvh_image_match.group(0)
                        html_content += f"""
                        <div class="section">
                            <h2>Biểu đồ DVH</h2>
                            <div class="img-container">
                                {dvh_image_tag}>
                            </div>
                        </div>
                        """
        
        # Thêm kết quả QA
        if "qa_results" in included_sections and qa_results:
            html_content += """
            <div class="section">
                <h2>Kết quả kiểm tra QA</h2>
            """
            
            all_passed = qa_results.get("all_passed", False)
            status = "Đạt" if all_passed else "Không đạt"
            status_class = "pass" if all_passed else "fail"
            
            html_content += f"""
                <p>Kết quả tổng thể: <span class="{status_class}">{status}</span></p>
                
                <h3>Chi tiết kiểm tra</h3>
                <table>
                    <tr>
                        <th>Kiểm tra</th>
                        <th>Kết quả</th>
                        <th>Chi tiết</th>
                    </tr>
            """
            
            if "results" in qa_results:
                for test_name, test_result in qa_results["results"].items():
                    passed = test_result.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    status_class = "pass" if passed else "fail"
                    comment = test_result.get("comment", "")
                    
                    html_content += f"""
                        <tr>
                            <td>{test_name}</td>
                            <td class="{status_class}">{status}</td>
                            <td>{comment}</td>
                        </tr>
                    """
            
            html_content += """
                </table>
            </div>
            """
        
        # Thêm kết quả KBP QA nếu có
        if "kbp_results" in included_sections and kbp_results:
            html_content += """
            <div class="section">
                <h2>Kết quả kiểm tra KBP QA</h2>
                <table>
                    <tr>
                        <th>Cấu trúc</th>
                        <th>Metric</th>
                        <th>Giá trị thực tế</th>
                        <th>Dự đoán</th>
                        <th>Sai lệch (%)</th>
                        <th>Kết quả</th>
                    </tr>
            """
            
            total_checks = len(kbp_results)
            passed_checks = sum(1 for result in kbp_results if isinstance(result, dict) and result.get("is_passed", False))
            pass_percent = (passed_checks / total_checks * 100) if total_checks > 0 else 0
            
            for result in kbp_results:
                if not isinstance(result, dict):
                    continue
                    
                structure_name = result.get("structure_name", "")
                metric_name = result.get("metric_name", "")
                actual_value = result.get("actual_value", 0)
                predicted_value = result.get("predicted_value", 0)
                delta_percent = result.get("delta_percent", 0)
                is_passed = result.get("is_passed", False)
                
                status = "Đạt" if is_passed else "Không đạt"
                status_class = "pass" if is_passed else "fail"
                
                html_content += f"""
                    <tr>
                        <td>{structure_name}</td>
                        <td>{metric_name}</td>
                        <td>{actual_value:.2f}</td>
                        <td>{predicted_value:.2f}</td>
                        <td>{delta_percent:.1f}%</td>
                        <td class="{status_class}">{status}</td>
                    </tr>
                """
            
            html_content += f"""
                </table>
                <p>Tổng số kiểm tra: {total_checks}</p>
                <p>Số kiểm tra đạt: {passed_checks} ({pass_percent:.1f}%)</p>
            </div>
            """
        
        # Thêm chỉ số tuân thủ và đồng nhất
        if ("conformity_indices" in included_sections or "homogeneity_indices" in included_sections) and "qa_results" in data:
            qa_results = data["qa_results"]
            if "results" in qa_results:
                if "conformity" in qa_results["results"] and "conformity_indices" in included_sections:
                    conformity = qa_results["results"]["conformity"]
                    html_content += """
                    <div class="section">
                        <h2>Chỉ số tuân thủ (Conformity)</h2>
                    """
                    
                    passed = conformity.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    status_class = "pass" if passed else "fail"
                    
                    html_content += f"""
                        <p>Kết quả: <span class="{status_class}">{status}</span></p>
                        <p>{conformity.get("comment", "")}</p>
                        
                        <table>
                            <tr>
                                <th>Chỉ số</th>
                                <th>Giá trị</th>
                            </tr>
                    """
                    
                    if "details" in conformity:
                        details = conformity["details"]
                        for key, value in details.items():
                            html_content += f"""
                                <tr>
                                    <td>{key}</td>
                                    <td>{value:.3f}</td>
                                </tr>
                            """
                    
                    html_content += """
                        </table>
                    </div>
                    """
                
                if "homogeneity" in qa_results["results"] and "homogeneity_indices" in included_sections:
                    homogeneity = qa_results["results"]["homogeneity"]
                    html_content += """
                    <div class="section">
                        <h2>Chỉ số đồng nhất (Homogeneity)</h2>
                    """
                    
                    passed = homogeneity.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    status_class = "pass" if passed else "fail"
                    
                    html_content += f"""
                        <p>Kết quả: <span class="{status_class}">{status}</span></p>
                        <p>{homogeneity.get("comment", "")}</p>
                        
                        <table>
                            <tr>
                                <th>Chỉ số</th>
                                <th>Giá trị</th>
                            </tr>
                    """
                    
                    if "details" in homogeneity:
                        details = homogeneity["details"]
                        for key, value in details.items():
                            html_content += f"""
                                <tr>
                                    <td>{key}</td>
                                    <td>{value:.3f}</td>
                                </tr>
                            """
                    
                    html_content += """
                        </table>
                    </div>
                    """
        
        # Thêm phần chân trang
        html_content += """
            <div class="footer">
                <p>Báo cáo được tạo bởi QuangStation V2 - Hệ thống lập kế hoạch xạ trị</p>
            </div>
        </body>
        </html>
        """
        
        # Ghi vào file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path
    
    except Exception as error:
        logger.error(f"Lỗi khi tạo báo cáo HTML tổng hợp: {str(error)}")
        raise


def _create_pdf_comprehensive_report(output_path: str, data: Dict[str, Any], temp_dir: str) -> str:
    """Tạo báo cáo tổng hợp PDF"""
    try:
        # Tạo báo cáo HTML trước
        html_path = os.path.join(temp_dir, "temp_report.html")
        _create_html_comprehensive_report(html_path, data, temp_dir)
        
        # Chuyển đổi HTML sang PDF
        try:
            import weasyprint
            weasyprint.HTML(html_path).write_pdf(output_path)
            return output_path
        except ImportError:
            logger.warning("Thư viện weasyprint không được cài đặt. Sử dụng báo cáo HTML thay thế.")
            # Sao chép file HTML sang vị trí đầu ra
            output_html_path = output_path.replace('.pdf', '.html')
            shutil.copy2(html_path, output_html_path)
            return output_html_path
    
    except Exception as error:
        logger.error(f"Lỗi khi tạo báo cáo PDF tổng hợp: {str(error)}")
        raise


def _create_docx_comprehensive_report(output_path: str, data: Dict[str, Any], temp_dir: str) -> str:
    """Tạo báo cáo tổng hợp DOCX"""
    try:
        # Kiểm tra thư viện python-docx
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("Thư viện python-docx không được cài đặt. Sử dụng báo cáo HTML thay thế.")
            # Tạo báo cáo HTML thay thế
            output_html_path = output_path.replace('.docx', '.html')
            return _create_html_comprehensive_report(output_html_path, data, temp_dir)
        
        # Thu thập dữ liệu
        patient_info = data.get("patient_info", {})
        plan_info = data.get("plan_info", {})
        dvh_data = data.get("dvh_data", {})
        dose_metrics = data.get("dose_metrics", {})
        qa_results = data.get("qa_results", {})
        kbp_results = data.get("kbp_results", [])
        included_sections = data.get("included_sections", [])
        
        # Tạo document mới
        document = Document()
        
        # Tiêu đề báo cáo
        title = document.add_heading(data.get("title", "Báo cáo tổng hợp kế hoạch xạ trị"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ngày tạo
        date_paragraph = document.add_paragraph(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm thông tin bệnh nhân
        if "patient_info" in included_sections:
            document.add_heading("Thông tin bệnh nhân", level=1)
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Tiêu đề bảng
            header_cells = table.rows[0].cells
            header_cells[0].text = "Thông tin"
            header_cells[1].text = "Giá trị"
            
            # Dữ liệu bảng
            for key, value in patient_info.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
            
            document.add_paragraph()
        
        # Thêm thông tin kế hoạch
        if "plan_info" in included_sections:
            document.add_heading("Thông tin kế hoạch", level=1)
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Tiêu đề bảng
            header_cells = table.rows[0].cells
            header_cells[0].text = "Thông tin"
            header_cells[1].text = "Giá trị"
            
            # Dữ liệu bảng
            for key, value in plan_info.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
            
            document.add_paragraph()
        
        # Thêm kết quả QA
        if "qa_results" in included_sections and qa_results:
            document.add_heading("Kết quả kiểm tra QA", level=1)
            
            all_passed = qa_results.get("all_passed", False)
            status = "Đạt" if all_passed else "Không đạt"
            
            result_paragraph = document.add_paragraph(f"Kết quả tổng thể: {status}")
            # Đặt màu cho kết quả
            for run in result_paragraph.runs:
                if status == "Đạt":
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Màu xanh lá
                else:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Màu đỏ
            
            document.add_heading("Chi tiết kiểm tra", level=2)
            table = document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Tiêu đề bảng
            header_cells = table.rows[0].cells
            header_cells[0].text = "Kiểm tra"
            header_cells[1].text = "Kết quả"
            header_cells[2].text = "Chi tiết"
            
            # Dữ liệu bảng
            if "results" in qa_results:
                for test_name, test_result in qa_results["results"].items():
                    passed = test_result.get("passed", False)
                    status = "Đạt" if passed else "Không đạt"
                    comment = test_result.get("comment", "")
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = test_name
                    row_cells[1].text = status
                    row_cells[2].text = comment
            
            document.add_paragraph()
        
        # Thêm kết quả KBP QA
        if "kbp_results" in included_sections and kbp_results:
            document.add_heading("Kết quả kiểm tra KBP QA", level=1)
            
            total_checks = len(kbp_results)
            passed_checks = sum(1 for result in kbp_results if isinstance(result, dict) and result.get("is_passed", False))
            pass_percent = (passed_checks / total_checks * 100) if total_checks > 0 else 0
            
            document.add_paragraph(f"Tổng số kiểm tra: {total_checks}")
            document.add_paragraph(f"Số kiểm tra đạt: {passed_checks} ({pass_percent:.1f}%)")
            
            table = document.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Tiêu đề bảng
            header_cells = table.rows[0].cells
            header_cells[0].text = "Cấu trúc"
            header_cells[1].text = "Metric"
            header_cells[2].text = "Giá trị thực tế"
            header_cells[3].text = "Dự đoán"
            header_cells[4].text = "Sai lệch (%)"
            header_cells[5].text = "Kết quả"
            
            # Dữ liệu bảng
            for result in kbp_results:
                if not isinstance(result, dict):
                    continue
                    
                structure_name = result.get("structure_name", "")
                metric_name = result.get("metric_name", "")
                actual_value = result.get("actual_value", 0)
                predicted_value = result.get("predicted_value", 0)
                delta_percent = result.get("delta_percent", 0)
                is_passed = result.get("is_passed", False)
                
                status = "Đạt" if is_passed else "Không đạt"
                
                row_cells = table.add_row().cells
                row_cells[0].text = structure_name
                row_cells[1].text = metric_name
                row_cells[2].text = f"{actual_value:.2f}"
                row_cells[3].text = f"{predicted_value:.2f}"
                row_cells[4].text = f"{delta_percent:.1f}%"
                row_cells[5].text = status
            
            document.add_paragraph()
        
        # Thêm chân trang
        footer_paragraph = document.add_paragraph("Báo cáo được tạo bởi QuangStation V2 - Hệ thống lập kế hoạch xạ trị")
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Lưu document
        document.save(output_path)
        
        return output_path
    
    except Exception as error:
        logger.error(f"Lỗi khi tạo báo cáo DOCX tổng hợp: {str(error)}")
        raise 