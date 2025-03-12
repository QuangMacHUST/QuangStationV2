"""
Module báo cáo điều trị cho hệ thống QuangStation V2
"""

import os
import json
from datetime import datetime
import matplotlib.pyplot as plt
import numpy as np
from typing import Dict, List, Tuple, Optional, Union, Any
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Image, Spacer, PageBreak
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from io import BytesIO

# Thêm thư viện mới cho DOCX và HTML
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import base64
import html

from quangstation.clinical.plan_evaluation.dvh import DVHCalculator
from quangstation.core.utils.logging import get_logger

logger = get_logger(__name__)

class TreatmentReport:
    """Lớp tạo báo cáo điều trị chi tiết"""
    
    def __init__(self, patient_data, plan_data, dose_data, structures):
        """
        Khởi tạo báo cáo điều trị
        
        Args:
            patient_data (dict): Thông tin bệnh nhân
            plan_data (dict): Thông tin kế hoạch điều trị
            dose_data (np.ndarray): Dữ liệu liều
            structures (dict): Thông tin các cấu trúc
        """
        self.patient_data = patient_data
        self.plan_data = plan_data
        self.dose_data = dose_data
        self.structures = structures
        
        # Tạo DVH Calculator
        self.dvh_calculator = DVHCalculator()
        self.dvh_calculator.set_dose_data(self.dose_data)
        
        # Tính toán DVH cho tất cả cấu trúc
        self.dvh_data = {}
        for name, structure in self.structures.items():
            self.dvh_calculator.add_structure(name, structure)
            self.dvh_data[name] = self.dvh_calculator.calculate_dvh(structure_name=name)
        
        # Tính các thống kê liều
        self._calculate_dose_statistics()
        
        logger.info("Khởi tạo báo cáo điều trị")
    
    def _calculate_dose_statistics(self):
        """Tính toán các thống kê liều cho báo cáo"""
        self.dose_stats = {
            "global": {
                "max_dose": float(np.max(self.dose_data)),
                "min_dose": float(np.min(self.dose_data)),
                "mean_dose": float(np.mean(self.dose_data))
            },
            "structures": {}
        }
        
        # Tính toán các thống kê cho từng cấu trúc
        for name, dvh_data in self.dvh_data.items():
            self.dose_stats["structures"][name] = {
                "min_dose": dvh_data.get("min_dose", 0),
                "max_dose": dvh_data.get("max_dose", 0),
                "mean_dose": dvh_data.get("mean_dose", 0),
                "d95": dvh_data.get("d95", 0),
                "d90": dvh_data.get("d90", 0),
                "d50": dvh_data.get("d50", 0),
                "v95": dvh_data.get("v95", 0),
                "v90": dvh_data.get("v90", 0),
                "v50": dvh_data.get("v50", 0),
                "volume_cc": dvh_data.get("volume_cc", 0),
                "conformity_index": dvh_data.get("conformity_index", 0),
                "homogeneity_index": dvh_data.get("homogeneity_index", 0)
            }
    
    def generate_pdf_report(self, output_path=None):
        """
        Tạo báo cáo PDF chi tiết
        
        Args:
            output_path (str): Đường dẫn file PDF
        """
        if output_path is None:
            output_path = f"treatment_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        try:
            # Tạo PDF
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Tạo style cho tiêu đề
            title_style = ParagraphStyle(
                'TitleStyle',
                parent=styles['Title'],
                fontSize=16,
                alignment=TA_CENTER,
                spaceAfter=20
            )
            
            # Tiêu đề báo cáo
            title = Paragraph("BÁO CÁO KẾ HOẠCH XẠ TRỊ", title_style)
            story.append(title)
            story.append(Spacer(1, 0.5*cm))
            
            # Thông tin bệnh nhân
            patient_info = [
                ["THÔNG TIN BỆNH NHÂN", ""],
                ["Họ và tên", self.patient_data.get('patient_name', 'N/A')],
                ["Mã bệnh nhân", self.patient_data.get('patient_id', 'N/A')],
                ["Ngày sinh", self.patient_data.get('birth_date', 'N/A')],
                ["Giới tính", self.patient_data.get('sex', 'N/A')],
                ["Chẩn đoán", self.patient_data.get('diagnosis', 'N/A')],
                ["Bác sĩ", self.patient_data.get('doctor', 'N/A')]
            ]
            patient_table = Table(patient_info, colWidths=[4*cm, 12*cm])
            patient_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (1,0), colors.darkblue),
                ('TEXTCOLOR', (0,0), (1,0), colors.white),
                ('SPAN', (0,0), (1,0)),
                ('ALIGN', (0,0), (-1,0), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 12),
                ('BOTTOMPADDING', (0,0), (-1,0), 12),
                ('BACKGROUND', (0,1), (0,-1), colors.lightgrey),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
            ]))
            story.append(patient_table)
            story.append(Spacer(1, 0.5*cm))
            
            # Thông tin kế hoạch
            plan_info = [
                ["THÔNG TIN KẾ HOẠCH", ""],
                ["Tên kế hoạch", self.plan_data.get('plan_name', 'N/A')],
                ["Kỹ thuật", self.plan_data.get('technique', 'N/A')],
                ["Liều tổng", f"{self.plan_data.get('total_dose', 'N/A')} Gy"],
                ["Số phân đoạn", self.plan_data.get('fractions', 'N/A')],
                ["Liều mỗi phân đoạn", f"{self.plan_data.get('fraction_dose', 'N/A')} Gy"],
                ["Thuật toán tính liều", self.plan_data.get('dose_algorithm', 'N/A')],
                ["Ngày tạo kế hoạch", self.plan_data.get('created_date', 'N/A')]
            ]
            plan_table = Table(plan_info, colWidths=[4*cm, 12*cm])
            plan_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (1,0), colors.darkblue),
                ('TEXTCOLOR', (0,0), (1,0), colors.white),
                ('SPAN', (0,0), (1,0)),
                ('ALIGN', (0,0), (-1,0), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 12),
                ('BOTTOMPADDING', (0,0), (-1,0), 12),
                ('BACKGROUND', (0,1), (0,-1), colors.lightgrey),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
            ]))
            story.append(plan_table)
            story.append(Spacer(1, 0.5*cm))
            
            # Thông tin chùm tia
            beam_data = self.plan_data.get('beams', [])
            if beam_data:
                story.append(Paragraph("THÔNG TIN CHÙM TIA", styles['Heading2']))
                beam_rows = [["ID", "Năng lượng", "Góc Gantry", "Góc Bàn", "Góc Collimator", "SSD", "MU"]]
                
                for beam in beam_data:
                    beam_rows.append([
                        beam.get('id', 'N/A'),
                        f"{beam.get('energy', 'N/A')} MV",
                        f"{beam.get('gantry_angle', 'N/A')}°",
                        f"{beam.get('couch_angle', 'N/A')}°",
                        f"{beam.get('collimator_angle', 'N/A')}°",
                        f"{beam.get('ssd', 'N/A')} cm",
                        f"{beam.get('mu', 'N/A')}"
                    ])
                
                beam_table = Table(beam_rows, colWidths=[2*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm])
                beam_table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.darkblue),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('GRID', (0,0), (-1,-1), 1, colors.black),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
                ]))
                story.append(beam_table)
                story.append(Spacer(1, 0.5*cm))
            
            # Trang mới cho phần DVH và thống kê liều
            story.append(PageBreak())
            story.append(Paragraph("BIỂU ĐỒ LIỀU - THỂ TÍCH (DVH)", styles['Heading2']))
            
            # Tạo DVH
            dvh_figure = self.dvh_calculator.create_figure(figsize=(8, 5))
            self.dvh_calculator.plot_dvh(figure=dvh_figure)
            
            # Lưu DVH làm ảnh tạm
            dvh_buffer = BytesIO()
            dvh_figure.savefig(dvh_buffer, format='png', dpi=300, bbox_inches='tight')
            dvh_image = Image(dvh_buffer)
            dvh_image.drawHeight = 10*cm
            dvh_image.drawWidth = 16*cm
            story.append(dvh_image)
            story.append(Spacer(1, 0.5*cm))
            
            # Bảng thống kê liều cho các cấu trúc
            story.append(Paragraph("THỐNG KÊ LIỀU CHO CÁC CẤU TRÚC", styles['Heading2']))
            
            dose_stat_rows = [["Cấu trúc", "Thể tích (cc)", "Liều Tối đa (Gy)", "Liều Trung bình (Gy)", "D95 (Gy)", "V95 (%)"]]
            
            for name, stats in self.dose_stats["structures"].items():
                dose_stat_rows.append([
                    name,
                    f"{stats['volume_cc']:.2f}",
                    f"{stats['max_dose']:.2f}",
                    f"{stats['mean_dose']:.2f}",
                    f"{stats['d95']:.2f}",
                    f"{stats['v95']:.2f}"
                ])
                
            dose_table = Table(dose_stat_rows, colWidths=[3*cm, 3*cm, 3*cm, 3*cm, 3*cm, 3*cm])
            dose_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.darkblue),
                ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
            ]))
            story.append(dose_table)
            
            # Thêm các chỉ số đánh giá kế hoạch nếu có
            if any('conformity_index' in stats for stats in self.dose_stats["structures"].values()):
                story.append(Spacer(1, 0.5*cm))
                story.append(Paragraph("CHỈ SỐ ĐÁNH GIÁ KẾ HOẠCH", styles['Heading2']))
                
                eval_rows = [["Cấu trúc", "Chỉ số Conformity", "Chỉ số Homogeneity"]]
                
                for name, stats in self.dose_stats["structures"].items():
                    if 'PTV' in name or 'TARGET' in name.upper():  # Chỉ xem xét PTV hoặc target
                        eval_rows.append([
                            name,
                            f"{stats.get('conformity_index', 'N/A')}",
                            f"{stats.get('homogeneity_index', 'N/A')}"
                        ])
                
                if len(eval_rows) > 1:  # Nếu có dữ liệu ngoài hàng tiêu đề
                    eval_table = Table(eval_rows, colWidths=[6*cm, 5*cm, 5*cm])
                    eval_table.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.darkblue),
                        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                        ('GRID', (0,0), (-1,-1), 1, colors.black),
                        ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
                    ]))
                    story.append(eval_table)
            
            # Xây dựng PDF
            doc.build(story)
            
            logger.info(f"Đã tạo báo cáo PDF tại: {output_path}")
            return output_path
        
        except Exception as error:
            logger.error(f"Lỗi khi tạo báo cáo PDF: {error}")
            return None
    
    def generate_docx_report(self, output_path=None):
        """
        Tạo báo cáo chi tiết dạng DOCX
        
        Args:
            output_path (str): Đường dẫn file DOCX
        """
        if output_path is None:
            output_path = f"treatment_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        try:
            # Tạo document
            document = Document()
            
            # Tiêu đề báo cáo
            title = document.add_heading('BÁO CÁO KẾ HOẠCH XẠ TRỊ', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thông tin bệnh nhân
            document.add_heading('THÔNG TIN BỆNH NHÂN', level=1)
            table = document.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            
            cells = [
                ("Họ và tên", self.patient_data.get('patient_name', 'N/A')),
                ("Mã bệnh nhân", self.patient_data.get('patient_id', 'N/A')),
                ("Ngày sinh", self.patient_data.get('birth_date', 'N/A')),
                ("Giới tính", self.patient_data.get('sex', 'N/A')),
                ("Chẩn đoán", self.patient_data.get('diagnosis', 'N/A')),
                ("Bác sĩ", self.patient_data.get('doctor', 'N/A'))
            ]
            
            for i, (key, value) in enumerate(cells):
                row = table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = value
            
            document.add_paragraph('')
            
            # Thông tin kế hoạch
            document.add_heading('THÔNG TIN KẾ HOẠCH', level=1)
            plan_table = document.add_table(rows=7, cols=2)
            plan_table.style = 'Table Grid'
            
            plan_cells = [
                ("Tên kế hoạch", self.plan_data.get('plan_name', 'N/A')),
                ("Kỹ thuật", self.plan_data.get('technique', 'N/A')),
                ("Liều tổng", f"{self.plan_data.get('total_dose', 'N/A')} Gy"),
                ("Số phân đoạn", str(self.plan_data.get('fractions', 'N/A'))),
                ("Liều mỗi phân đoạn", f"{self.plan_data.get('fraction_dose', 'N/A')} Gy"),
                ("Thuật toán tính liều", self.plan_data.get('dose_algorithm', 'N/A')),
                ("Ngày tạo kế hoạch", self.plan_data.get('created_date', 'N/A'))
            ]
            
            for i, (key, value) in enumerate(plan_cells):
                row = plan_table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = value
            
            document.add_paragraph('')
            
            # Thông tin chùm tia
            beam_data = self.plan_data.get('beams', [])
            if beam_data:
                document.add_heading('THÔNG TIN CHÙM TIA', level=1)
                beam_table = document.add_table(rows=1, cols=7)
                beam_table.style = 'Table Grid'
                
                # Tiêu đề
                header_cells = beam_table.rows[0].cells
                header_cells[0].text = "ID"
                header_cells[1].text = "Năng lượng"
                header_cells[2].text = "Góc Gantry"
                header_cells[3].text = "Góc Bàn"
                header_cells[4].text = "Góc Collimator"
                header_cells[5].text = "SSD"
                header_cells[6].text = "MU"
                
                # Dữ liệu chùm tia
                for beam in beam_data:
                    row_cells = beam_table.add_row().cells
                    row_cells[0].text = beam.get('id', 'N/A')
                    row_cells[1].text = f"{beam.get('energy', 'N/A')} MV"
                    row_cells[2].text = f"{beam.get('gantry_angle', 'N/A')}°"
                    row_cells[3].text = f"{beam.get('couch_angle', 'N/A')}°"
                    row_cells[4].text = f"{beam.get('collimator_angle', 'N/A')}°"
                    row_cells[5].text = f"{beam.get('ssd', 'N/A')} cm"
                    row_cells[6].text = f"{beam.get('mu', 'N/A')}"
                
                document.add_paragraph('')
            
            # Tạo DVH
            document.add_heading('BIỂU ĐỒ LIỀU - THỂ TÍCH (DVH)', level=1)
            dvh_figure = self.dvh_calculator.create_figure(figsize=(8, 5))
            self.dvh_calculator.plot_dvh(figure=dvh_figure)
            
            # Lưu DVH làm ảnh tạm
            dvh_path = f"temp_dvh_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            dvh_figure.savefig(dvh_path, format='png', dpi=300, bbox_inches='tight')
            document.add_picture(dvh_path, width=Inches(6))
            
            # Xóa file tạm sau khi đã thêm vào báo cáo
            try:
                os.remove(dvh_path)
            except:
                pass
            
            document.add_paragraph('')
            
            # Bảng thống kê liều cho các cấu trúc
            document.add_heading('THỐNG KÊ LIỀU CHO CÁC CẤU TRÚC', level=1)
            dose_table = document.add_table(rows=1, cols=6)
            dose_table.style = 'Table Grid'
            
            # Tiêu đề
            header_cells = dose_table.rows[0].cells
            header_cells[0].text = "Cấu trúc"
            header_cells[1].text = "Thể tích (cc)"
            header_cells[2].text = "Liều Tối đa (Gy)"
            header_cells[3].text = "Liều Trung bình (Gy)"
            header_cells[4].text = "D95 (Gy)"
            header_cells[5].text = "V95 (%)"
            
            # Dữ liệu liều
            for name, stats in self.dose_stats["structures"].items():
                row_cells = dose_table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = f"{stats['volume_cc']:.2f}"
                row_cells[2].text = f"{stats['max_dose']:.2f}"
                row_cells[3].text = f"{stats['mean_dose']:.2f}"
                row_cells[4].text = f"{stats['d95']:.2f}"
                row_cells[5].text = f"{stats['v95']:.2f}"
            
            # Lưu văn bản
            document.save(output_path)
            
            logger.info(f"Đã tạo báo cáo DOCX tại: {output_path}")
            return output_path
        
        except Exception as error:
            logger.error(f"Lỗi khi tạo báo cáo DOCX: {error}")
            return None
    
    def generate_html_report(self, output_path=None):
        """
        Tạo báo cáo HTML
        
        Args:
            output_path (str): Đường dẫn file HTML
        """
        if output_path is None:
            output_path = f"treatment_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        
        try:
            # Tạo DVH
            dvh_figure = self.dvh_calculator.create_figure(figsize=(8, 5))
            self.dvh_calculator.plot_dvh(figure=dvh_figure)
            
            # Lưu DVH làm ảnh Base64
            dvh_buffer = BytesIO()
            dvh_figure.savefig(dvh_buffer, format='png', dpi=150, bbox_inches='tight')
            dvh_base64 = base64.b64encode(dvh_buffer.getvalue()).decode('utf-8')
            
            # Tạo HTML
            html_content = f"""
            <!DOCTYPE html>
            <html lang="vi">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Báo Cáo Kế Hoạch Xạ Trị</title>
                <style>
                    body {{ 
                        font-family: Arial, sans-serif; 
                        line-height: 1.6; 
                        margin: 20px;
                        color: #333;
                    }}
                    h1, h2 {{ 
                        color: #005cb9; 
                        text-align: center; 
                    }}
                    .container {{ 
                        max-width: 1000px; 
                        margin: 0 auto; 
                    }}
                    table {{ 
                        width: 100%; 
                        border-collapse: collapse; 
                        margin: 20px 0; 
                    }}
                    th, td {{ 
                        padding: 12px; 
                        text-align: left; 
                        border: 1px solid #ddd; 
                    }}
                    th {{ 
                        background-color: #005cb9; 
                        color: white; 
                    }}
                    tr:nth-child(even) {{ 
                        background-color: #f2f2f2; 
                    }}
                    .dvh-image {{ 
                        display: block;
                        margin: 20px auto;
                        max-width: 800px;
                    }}
                    .section {{ 
                        margin-bottom: 40px; 
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <h1>BÁO CÁO KẾ HOẠCH XẠ TRỊ</h1>
                    
                    <div class="section">
                        <h2>Thông Tin Bệnh Nhân</h2>
                        <table>
                            <tr>
                                <th>Thông Tin</th>
                                <th>Chi Tiết</th>
                            </tr>
                            <tr>
                                <td>Họ và tên</td>
                                <td>{html.escape(str(self.patient_data.get('patient_name', 'N/A')))}</td>
                            </tr>
                            <tr>
                                <td>Mã bệnh nhân</td>
                                <td>{html.escape(str(self.patient_data.get('patient_id', 'N/A')))}</td>
                            </tr>
                            <tr>
                                <td>Ngày sinh</td>
                                <td>{html.escape(str(self.patient_data.get('birth_date', 'N/A')))}</td>
                            </tr>
                            <tr>
                                <td>Giới tính</td>
                                <td>{html.escape(str(self.patient_data.get('sex', 'N/A')))}</td>
                            </tr>
                            <tr>
                                <td>Chẩn đoán</td>
                                <td>{html.escape(str(self.patient_data.get('diagnosis', 'N/A')))}</td>
                            </tr>
                            <tr>
                                <td>Bác sĩ</td>
                                <td>{html.escape(str(self.patient_data.get('doctor', 'N/A')))}</td>
                            </tr>
                        </table>
                    </div>
                    
                    <div class="section">
                        <h2>Thông Tin Kế Hoạch</h2>
                        <table>
                            <tr>
                                <th>Thông Tin</th>
                                <th>Chi Tiết</th>
                            </tr>
                            <tr>
                                <td>Tên kế hoạch</td>
                                <td>{html.escape(str(self.plan_data.get('plan_name', 'N/A')))}</td>
                            </tr>
                            <tr>
                                <td>Kỹ thuật</td>
                                <td>{html.escape(str(self.plan_data.get('technique', 'N/A')))}</td>
                            </tr>
                            <tr>
                                <td>Liều tổng</td>
                                <td>{html.escape(str(self.plan_data.get('total_dose', 'N/A')))} Gy</td>
                            </tr>
                            <tr>
                                <td>Số phân đoạn</td>
                                <td>{html.escape(str(self.plan_data.get('fractions', 'N/A')))}</td>
                            </tr>
                            <tr>
                                <td>Liều mỗi phân đoạn</td>
                                <td>{html.escape(str(self.plan_data.get('fraction_dose', 'N/A')))} Gy</td>
                            </tr>
                            <tr>
                                <td>Thuật toán tính liều</td>
                                <td>{html.escape(str(self.plan_data.get('dose_algorithm', 'N/A')))}</td>
                            </tr>
                            <tr>
                                <td>Ngày tạo kế hoạch</td>
                                <td>{html.escape(str(self.plan_data.get('created_date', 'N/A')))}</td>
                            </tr>
                        </table>
                    </div>
                    
                    <div class="section">
                        <h2>Biểu Đồ Liều - Thể Tích (DVH)</h2>
                        <img src="data:image/png;base64,{dvh_base64}" class="dvh-image" alt="Biểu đồ DVH">
                    </div>
                    
                    <div class="section">
                        <h2>Thống Kê Liều Cho Các Cấu Trúc</h2>
                        <table>
                            <tr>
                                <th>Cấu trúc</th>
                                <th>Thể tích (cc)</th>
                                <th>Liều Tối đa (Gy)</th>
                                <th>Liều Trung bình (Gy)</th>
                                <th>D95 (Gy)</th>
                                <th>V95 (%)</th>
                            </tr>
            """
            
            # Thêm dữ liệu thống kê liều
            for name, stats in self.dose_stats["structures"].items():
                html_content += f"""
                            <tr>
                                <td>{html.escape(name)}</td>
                                <td>{stats['volume_cc']:.2f}</td>
                                <td>{stats['max_dose']:.2f}</td>
                                <td>{stats['mean_dose']:.2f}</td>
                                <td>{stats['d95']:.2f}</td>
                                <td>{stats['v95']:.2f}</td>
                            </tr>
                """
            
            # Đóng bảng và thẻ HTML
            html_content += """
                        </table>
                    </div>
                </div>
            </body>
            </html>
            """
            
            # Lưu file HTML
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"Đã tạo báo cáo HTML tại: {output_path}")
            return output_path
        
        except Exception as error:
            logger.error(f"Lỗi khi tạo báo cáo HTML: {error}")
            return None
    
    def export_json_summary(self, output_path=None):
        """
        Xuất báo cáo dưới dạng JSON
        
        Args:
            output_path (str): Đường dẫn file JSON
        """
        if output_path is None:
            output_path = f"treatment_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        summary = {
            "patient_info": self.patient_data,
            "plan_details": self.plan_data,
            "dose_statistics": self.dose_stats,
            "created_at": datetime.now().isoformat(),
            "quangstation_version": "2.0.0"
        }
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(summary, f, indent=4, ensure_ascii=False)
            
            logger.info(f"Đã xuất báo cáo JSON tại: {output_path}")
            return output_path
        
        except Exception as error:
            logger.error(f"Lỗi khi xuất báo cáo JSON: {error}")
            return None

def generate_treatment_report(patient_data, plan_data, dose_data, structures, 
                              pdf_output=None, docx_output=None, html_output=None, json_output=None):
    """
    Hàm tiện ích để tạo báo cáo điều trị
    
    Args:
        patient_data (dict): Thông tin bệnh nhân
        plan_data (dict): Thông tin kế hoạch
        dose_data (np.ndarray): Dữ liệu liều
        structures (dict): Thông tin các cấu trúc
        pdf_output (str, optional): Đường dẫn file PDF
        docx_output (str, optional): Đường dẫn file DOCX
        html_output (str, optional): Đường dẫn file HTML
        json_output (str, optional): Đường dẫn file JSON
    
    Returns:
        dict: Các đường dẫn file đã tạo
    """
    report = TreatmentReport(patient_data, plan_data, dose_data, structures)
    
    results = {}
    
    if pdf_output is not None or docx_output is None and html_output is None and json_output is None:
        # Nếu không chỉ định output nào, mặc định tạo PDF
        results['pdf'] = report.generate_pdf_report(pdf_output)
    else:
        if pdf_output is not None:
            results['pdf'] = report.generate_pdf_report(pdf_output)
    
    if docx_output is not None:
        results['docx'] = report.generate_docx_report(docx_output)
    
    if html_output is not None:
        results['html'] = report.generate_html_report(html_output)
    
    if json_output is not None:
        results['json'] = report.export_json_summary(json_output)
    
    return results
